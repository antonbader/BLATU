# -*- coding: utf-8 -*-
"""
Dienstprogramm zur Erstellung von Word-Dokumenten aus Vorlagen
"""

import os
import warnings
from docx import Document
import re

# Monkey-patch docxcompose to handle missing templates in frozen environments
# This must be done BEFORE importing Composer or instantiating it,
# but since docxcompose imports properties at the top level, we need to patch the class method.

# XML content from docxcompose/templates/custom.xml
CUSTOM_XML_CONTENT = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"></Properties>"""

def patched_part_template(self):
    """Patched method to return hardcoded custom.xml content."""
    return CUSTOM_XML_CONTENT

# Warnung von docxcompose unterdrücken und Patch anwenden
with warnings.catch_warnings():
    warnings.filterwarnings("ignore", category=UserWarning, module='docxcompose.properties')
    from docxcompose import properties
    # Apply patch
    properties.CustomProperties._part_template = patched_part_template
    from docxcompose.composer import Composer


class WordGenerator:
    """Erstellt Word-Dokumente aus einer Vorlage."""

    def __init__(self, template_path):
        self.template_path = template_path

    def generate_certificate(self, output_path, data):
        """
        Erstellt eine einzelne Urkunde.
        :param output_path: Der Speicherpfad für die neue Word-Datei.
        :param data: Ein Dictionary mit den Platzhalter-Daten.
        """
        try:
            document = Document(self.template_path)
            self._replace_placeholders_in_document(document, data)
            document.save(output_path)
            return True, None
        except Exception as e:
            return False, str(e)

    def generate_schiesszettel(self, output_path, shooters_list, template_path, turnier_data):
        """
        Erstellt Schießzettel für eine Gruppe von Schützen in einer Datei.
        Verwendet docxcompose, um mehrere Seiten (eine pro 4 Schützen) zusammenzufügen.
        """
        try:
            # Sortieren nach Scheibe
            shooters_list.sort(key=lambda s: s.get('scheibe', 0))

            # Chunking in 4er Gruppen
            chunks = [shooters_list[i:i + 4] for i in range(0, len(shooters_list), 4)]

            master_doc = None
            composer = None

            for chunk in chunks:
                # Vorbereiten der Daten für diesen Chunk
                data = {}

                # Turnierdaten
                data["[Turniername]"] = turnier_data.get('name', '')
                data["[Turnierdatum]"] = turnier_data.get('datum', '')

                # Schützendaten
                for i in range(4):
                    suffix = f"_{i+1}"
                    if i < len(chunk):
                        schuetze = chunk[i]
                        data[f"[Name{suffix}]"] = str(schuetze.get('name', ''))
                        data[f"[Vorname{suffix}]"] = str(schuetze.get('vorname', ''))
                        data[f"[Gruppe{suffix}]"] = str(schuetze.get('gruppe', ''))
                        data[f"[Scheibe{suffix}]"] = str(schuetze.get('scheibe', ''))
                    else:
                        # Leere Platzhalter für fehlende Schützen
                        data[f"[Name{suffix}]"] = ""
                        data[f"[Vorname{suffix}]"] = ""
                        data[f"[Gruppe{suffix}]"] = ""
                        data[f"[Scheibe{suffix}]"] = ""

                # Dokument für diesen Chunk erstellen
                doc = Document(template_path)
                self._replace_placeholders_in_document(doc, data)

                # Mergen
                if master_doc is None:
                    master_doc = doc
                    composer = Composer(master_doc)
                else:
                    # Seiteumbruch im Master-Dokument hinzufügen, BEVOR das neue Dokument angehängt wird
                    master_doc.add_page_break()
                    composer.append(doc)

            if master_doc:
                composer.save(output_path)
                return True, None
            else:
                 # Fallback: Leeres Dokument oder Fehler, sollte nicht passieren wenn Liste nicht leer
                 return False, "Keine Schützen für Schießzettel vorhanden."

        except Exception as e:
            return False, str(e)

    def _replace_placeholders_in_document(self, document, data):
        """Ersetzt Platzhalter im gesamten Dokument (Paragraphen und Tabellen)."""
        # Paragraphs
        for p in document.paragraphs:
            for key, value in data.items():
                self._replace_placeholder_in_paragraph(p, key, str(value))

        # Tables
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for key, value in data.items():
                            self._replace_placeholder_in_paragraph(p, key, str(value))

    def _replace_placeholder_in_paragraph(self, paragraph, placeholder, value):
        """
        Replaces a placeholder in a paragraph, handling split runs.
        Optimized to preserve formatting.
        """
        # Case A: Check if placeholder exists in a single run.
        # This is the optimal case as it preserves specific run formatting perfectly.
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)
                return

        # Case B: Placeholder is split across multiple runs.
        # We need to rebuild the paragraph text, but we try to preserve the style
        # of the FIRST run in the paragraph, as requested.
        full_text = ''.join(run.text for run in paragraph.runs)

        if placeholder in full_text:
            # Capture style from the first run before clearing
            first_run_style = {}
            if paragraph.runs:
                first_run = paragraph.runs[0]
                first_run_style = {
                    'bold': first_run.bold,
                    'italic': first_run.italic,
                    'underline': first_run.underline,
                    'strike': first_run.font.strike,
                    'size': first_run.font.size,
                    'name': first_run.font.name,
                    'color_rgb': first_run.font.color.rgb,
                    # We can add more if needed
                }

            # Clear existing paragraph content
            for run in paragraph.runs:
                run.text = ''

            # Replace and add new content
            new_text = full_text.replace(placeholder, value)

            # Add a new run with the new text
            new_run = paragraph.add_run(new_text)

            # Apply captured style to the new run
            if first_run_style:
                new_run.bold = first_run_style.get('bold')
                new_run.italic = first_run_style.get('italic')
                new_run.underline = first_run_style.get('underline')

                if first_run_style.get('strike') is not None:
                    new_run.font.strike = first_run_style.get('strike')

                if first_run_style.get('size') is not None:
                    new_run.font.size = first_run_style.get('size')

                if first_run_style.get('name') is not None:
                    new_run.font.name = first_run_style.get('name')

                if first_run_style.get('color_rgb') is not None:
                    new_run.font.color.rgb = first_run_style.get('color_rgb')
